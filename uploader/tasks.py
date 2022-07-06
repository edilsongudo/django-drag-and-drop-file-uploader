from io import BytesIO
from os import remove
from os.path import splitext
from time import sleep

import requests
from celery import shared_task
from django.conf import settings
from docx import Document
from docx.shared import Inches
from storages.backends.s3boto3 import S3Boto3Storage

from .utils import milliseconds_to_hours


def upload_file(file):
    def read_file(file, chunk_size=5242880):
        with open(file, 'rb') as _file:
            while True:
                data = _file.read(chunk_size)
                if not data:
                    break
                yield data

    # Step 0: Upload file
    response = requests.post(
        'https://api.assemblyai.com/v2/upload',
        headers=headers,
        data=read_file(file),
    )
    upload_url = response.json()['upload_url']
    remove(file)
    return upload_url


@shared_task
def transcribe(upload_url, filename, file=None):

    headers = {
        'authorization': settings.ASSEMBLYAI_TOKEN,
        'content-type': 'application/json',
    }

    if file:
        upload_url = upload_file()

    # Step 1: Submit file for transcription
    endpoint = 'https://api.assemblyai.com/v2/transcript'
    json_data = {'audio_url': upload_url, 'speaker_labels': True}
    response = requests.post(endpoint, json=json_data, headers=headers)
    polling_endpoint = endpoint + '/' + response.json()['id']

    # Step 2: Poll Assembly Ai
    status = 'submited'
    while status not in ('completed', 'error'):
        response = requests.get(polling_endpoint, headers=headers)
        status = response.json()['status']
        sleep(5)

    if status == 'error':
        raise Exception('Transcription failed')

    # Step 3: Create word document and return the url
    utterances = response.json()['utterances']
    document = Document()

    for utterance in utterances:
        start = milliseconds_to_hours(utterance['start'])
        document.add_paragraph(f'{utterance["speaker"]} at {start}')
        document.add_paragraph(f'{utterance["text"]}')
        document.add_paragraph()

    with BytesIO() as fileobj:
        document.save(fileobj)
        fileobj.seek(0)

        storage = S3Boto3Storage()
        file_name = splitext(filename)[0] + '.docx'
        filename = storage.save(file_name, fileobj)
        file_url = storage.url(filename)

    return file_url
